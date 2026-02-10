#!/usr/bin/env python3
"""
Quick test to verify confidence indicators are removed from generated documents.
"""

import os
import sys
from docx import Document

# Test output file from previous runs
test_doc_path = "backend/test_output/sample_results/phase1_e2e_result.docx"

# Alternative: check if there's a recent document in output/
output_docs = []
if os.path.exists("backend/output"):
    for file in os.listdir("backend/output"):
        if file.endswith(".docx"):
            full_path = os.path.join("backend/output", file)
            output_docs.append(full_path)

# Use most recent document
if output_docs:
    test_doc_path = max(output_docs, key=os.path.getmtime)
elif os.path.exists(test_doc_path):
    pass
else:
    print("❌ No test document found to verify")
    print("Please run: cd backend && python tests/e2e/test_phase1_e2e.py")
    sys.exit(1)

print(f"📄 Checking document: {test_doc_path}")
print()

# Load the document
doc = Document(test_doc_path)

# Forbidden phrases (should NOT appear in document)
forbidden_phrases = [
    "Quality Indicator:",
    "Confidence:",
    "Average Confidence:",
    "High Confidence Steps",
    "confidence score",
    "Very High",
    "Very Low"
]

# Search through all paragraphs
found_issues = []

for i, para in enumerate(doc.paragraphs):
    text = para.text
    for phrase in forbidden_phrases:
        if phrase in text:
            found_issues.append({
                'paragraph': i,
                'phrase': phrase,
                'context': text[:100] + ('...' if len(text) > 100 else '')
            })

# Search through tables
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text
            for phrase in forbidden_phrases:
                if phrase in text:
                    found_issues.append({
                        'table': table_idx,
                        'row': row_idx,
                        'cell': cell_idx,
                        'phrase': phrase,
                        'context': text[:100] + ('...' if len(text) > 100 else '')
                    })

# Report results
print("=" * 70)
print("CONFIDENCE INDICATOR REMOVAL TEST")
print("=" * 70)
print()

if found_issues:
    print(f"❌ FAILED: Found {len(found_issues)} confidence indicators in document:")
    print()
    for issue in found_issues:
        if 'paragraph' in issue:
            print(f"  Paragraph {issue['paragraph']}:")
            print(f"    Phrase: '{issue['phrase']}'")
            print(f"    Context: {issue['context']}")
        else:
            print(f"  Table {issue['table']}, Row {issue['row']}, Cell {issue['cell']}:")
            print(f"    Phrase: '{issue['phrase']}'")
            print(f"    Context: {issue['context']}")
        print()
    sys.exit(1)
else:
    print("✅ PASSED: No confidence indicators found in document!")
    print()
    print("Verified that these phrases are NOT in the document:")
    for phrase in forbidden_phrases:
        print(f"  ✓ '{phrase}'")
    print()
    print("Document is clean and ready for production! 🎉")
    sys.exit(0)
