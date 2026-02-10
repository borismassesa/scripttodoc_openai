"""
Unit tests for document conversion service.
Tests PDF, PPTX, and DOCX conversion functionality.
"""

import pytest
from pathlib import Path
from docx import Document

from script_to_doc.converters.conversion_service import (
    ConversionService,
    DocumentFormat,
    get_conversion_service
)
from script_to_doc.converters.base import ConversionError


class TestConversionService:
    """Test suite for ConversionService."""

    def test_service_initialization(self):
        """Test that conversion service initializes correctly."""
        service = ConversionService()
        assert service is not None
        assert hasattr(service, '_converters')
        assert len(service._converters) == 2  # PDF and PPTX

    def test_singleton_pattern(self):
        """Test that get_conversion_service returns same instance."""
        service1 = get_conversion_service()
        service2 = get_conversion_service()
        assert service1 is service2

    def test_supported_formats(self):
        """Test that all expected formats are supported."""
        service = ConversionService()
        formats = service.get_supported_formats()

        assert DocumentFormat.DOCX in formats
        assert DocumentFormat.PDF in formats
        assert DocumentFormat.PPTX in formats
        assert len(formats) == 3

    def test_is_format_supported(self):
        """Test format support checking."""
        service = ConversionService()

        assert service.is_format_supported(DocumentFormat.DOCX) is True
        assert service.is_format_supported(DocumentFormat.PDF) is True
        assert service.is_format_supported(DocumentFormat.PPTX) is True


class TestDOCXConversion:
    """Test suite for DOCX handling (no conversion needed)."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file for testing."""
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document with some content.')
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('Section 1 content here.')

        docx_path = tmp_path / "test.docx"
        doc.save(docx_path)
        return docx_path

    def test_docx_no_conversion(self, sample_docx):
        """Test that DOCX format returns original file without conversion."""
        service = ConversionService()

        result = service.convert_document(
            input_path=sample_docx,
            output_format=DocumentFormat.DOCX
        )

        assert result == sample_docx
        assert result.exists()


class TestPDFConversion:
    """Test suite for PDF conversion."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file for testing."""
        doc = Document()
        doc.add_heading('PDF Test Document', 0)
        doc.add_paragraph('This document will be converted to PDF.')
        doc.add_paragraph('It contains multiple paragraphs.')

        docx_path = tmp_path / "test.docx"
        doc.save(docx_path)
        return docx_path

    def test_pdf_conversion_success(self, sample_docx, tmp_path):
        """Test successful DOCX to PDF conversion."""
        service = ConversionService()
        output_path = tmp_path / "output.pdf"

        result = service.convert_document(
            input_path=sample_docx,
            output_format=DocumentFormat.PDF,
            output_path=output_path
        )

        assert result.exists()
        assert result.suffix == '.pdf'
        assert result.stat().st_size > 0
        assert result == output_path

    def test_pdf_conversion_auto_output_path(self, sample_docx):
        """Test PDF conversion with automatic output path."""
        service = ConversionService()

        result = service.convert_document(
            input_path=sample_docx,
            output_format=DocumentFormat.PDF
        )

        assert result.exists()
        assert result.suffix == '.pdf'
        assert result.parent == sample_docx.parent
        assert result.stem == sample_docx.stem

    def test_pdf_conversion_missing_input(self, tmp_path):
        """Test error handling when input file doesn't exist."""
        service = ConversionService()
        fake_path = tmp_path / "nonexistent.docx"

        with pytest.raises(ConversionError) as exc_info:
            service.convert_document(
                input_path=fake_path,
                output_format=DocumentFormat.PDF
            )

        assert "Input file not found" in str(exc_info.value)


class TestPPTXConversion:
    """Test suite for PPTX conversion."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file for testing."""
        doc = Document()
        doc.add_heading('PowerPoint Test', 0)
        doc.add_paragraph('This will become a PowerPoint presentation.')
        doc.add_heading('Slide 1', level=1)
        doc.add_paragraph('Content for slide 1.')

        docx_path = tmp_path / "test.docx"
        doc.save(docx_path)
        return docx_path

    @pytest.mark.skip(reason="LibreOffice PPTX conversion requires additional setup")
    def test_pptx_conversion_success(self, sample_docx, tmp_path):
        """Test successful DOCX to PPTX conversion."""
        service = ConversionService()
        output_path = tmp_path / "output.pptx"

        result = service.convert_document(
            input_path=sample_docx,
            output_format=DocumentFormat.PPTX,
            output_path=output_path
        )

        assert result.exists()
        assert result.suffix == '.pptx'
        assert result.stat().st_size > 0
        assert result == output_path

    @pytest.mark.skip(reason="LibreOffice PPTX conversion requires additional setup")
    def test_pptx_conversion_auto_output_path(self, sample_docx):
        """Test PPTX conversion with automatic output path."""
        service = ConversionService()

        result = service.convert_document(
            input_path=sample_docx,
            output_format=DocumentFormat.PPTX
        )

        assert result.exists()
        assert result.suffix == '.pptx'
        assert result.parent == sample_docx.parent

    def test_pptx_conversion_missing_input(self, tmp_path):
        """Test error handling when input file doesn't exist."""
        service = ConversionService()
        fake_path = tmp_path / "nonexistent.docx"

        with pytest.raises(ConversionError) as exc_info:
            service.convert_document(
                input_path=fake_path,
                output_format=DocumentFormat.PPTX
            )

        assert "Input file not found" in str(exc_info.value)


class TestErrorHandling:
    """Test suite for error handling scenarios."""

    def test_invalid_format_type(self, tmp_path):
        """Test error when invalid format type is passed."""
        service = ConversionService()

        doc = Document()
        doc.add_paragraph("Test")
        docx_path = tmp_path / "test.docx"
        doc.save(docx_path)

        with pytest.raises((ValueError, AttributeError)):
            service.convert_document(
                input_path=docx_path,
                output_format="invalid_format"  # type: ignore
            )

    def test_conversion_error_propagation(self, tmp_path):
        """Test that ConversionError is properly propagated."""
        service = ConversionService()
        fake_path = tmp_path / "nonexistent.docx"

        with pytest.raises(ConversionError):
            service.convert_document(
                input_path=fake_path,
                output_format=DocumentFormat.PDF
            )


class TestIntegration:
    """Integration tests for full conversion workflow."""

    @pytest.fixture
    def complex_docx(self, tmp_path):
        """Create a more complex DOCX for testing."""
        doc = Document()

        # Title
        doc.add_heading('Training Document', 0)

        # Add multiple sections
        for i in range(3):
            doc.add_heading(f'Section {i+1}', level=1)
            doc.add_paragraph(f'This is content for section {i+1}.')
            doc.add_paragraph('Another paragraph with more text.')

            # Add a subsection
            doc.add_heading(f'Subsection {i+1}.1', level=2)
            doc.add_paragraph('Subsection content here.')

        # Add a list
        doc.add_paragraph('Important points:', style='List Bullet')
        for j in range(5):
            doc.add_paragraph(f'Point {j+1}', style='List Bullet')

        docx_path = tmp_path / "complex_test.docx"
        doc.save(docx_path)
        return docx_path

    @pytest.mark.skip(reason="LibreOffice PPTX conversion requires additional setup")
    def test_convert_to_all_formats(self, complex_docx, tmp_path):
        """Test converting a document to all supported formats."""
        service = ConversionService()
        results = {}

        # Convert to PDF
        pdf_path = tmp_path / "output.pdf"
        results['pdf'] = service.convert_document(
            input_path=complex_docx,
            output_format=DocumentFormat.PDF,
            output_path=pdf_path
        )

        # Convert to PPTX
        pptx_path = tmp_path / "output.pptx"
        results['pptx'] = service.convert_document(
            input_path=complex_docx,
            output_format=DocumentFormat.PPTX,
            output_path=pptx_path
        )

        # "Convert" to DOCX (no conversion)
        results['docx'] = service.convert_document(
            input_path=complex_docx,
            output_format=DocumentFormat.DOCX
        )

        # Verify all conversions succeeded
        assert results['pdf'].exists() and results['pdf'].suffix == '.pdf'
        assert results['pptx'].exists() and results['pptx'].suffix == '.pptx'
        assert results['docx'].exists() and results['docx'].suffix == '.docx'

        # Verify file sizes are reasonable
        assert results['pdf'].stat().st_size > 1000  # At least 1KB
        assert results['pptx'].stat().st_size > 1000
        assert results['docx'].stat().st_size > 1000

    def test_sequential_conversions(self, complex_docx, tmp_path):
        """Test multiple sequential conversions."""
        service = ConversionService()

        # Convert to PDF twice
        pdf1 = service.convert_document(
            input_path=complex_docx,
            output_format=DocumentFormat.PDF,
            output_path=tmp_path / "output1.pdf"
        )

        pdf2 = service.convert_document(
            input_path=complex_docx,
            output_format=DocumentFormat.PDF,
            output_path=tmp_path / "output2.pdf"
        )

        assert pdf1.exists() and pdf2.exists()
        assert pdf1 != pdf2

        # File sizes should be similar (within 10%)
        size1 = pdf1.stat().st_size
        size2 = pdf2.stat().st_size
        assert abs(size1 - size2) / max(size1, size2) < 0.1


# Performance tests (marked as slow)
class TestPerformance:
    """Performance tests for conversion operations."""

    @pytest.mark.slow
    def test_pdf_conversion_performance(self, tmp_path):
        """Test that PDF conversion completes within reasonable time."""
        import time

        # Create a moderately sized document
        doc = Document()
        doc.add_heading('Performance Test Document', 0)

        for i in range(20):  # 20 pages worth of content
            doc.add_heading(f'Section {i+1}', level=1)
            for j in range(10):
                doc.add_paragraph('Lorem ipsum dolor sit amet. ' * 20)

        docx_path = tmp_path / "perf_test.docx"
        doc.save(docx_path)

        service = ConversionService()

        start_time = time.time()
        result = service.convert_document(
            input_path=docx_path,
            output_format=DocumentFormat.PDF,
            output_path=tmp_path / "perf_output.pdf"
        )
        end_time = time.time()

        duration = end_time - start_time

        assert result.exists()
        assert duration < 30  # Should complete within 30 seconds
        print(f"\nPDF conversion took {duration:.2f} seconds")


if __name__ == "__main__":
    # Run tests with pytest
    pytest.main([__file__, "-v", "--tb=short"])
